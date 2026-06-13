import os
import sqlite3
import zipfile
import tarfile
import gzip
import shutil
import subprocess
import hashlib
import secrets
from pathlib import Path
from datetime import datetime, timedelta
import tkinter as tk
from tkinter import filedialog, messagebox, ttk


# ============================================================
# Optional imports
# ============================================================

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None
    ImageDraw = None
    ImageFont = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except ImportError:
    canvas = None
    letter = None

try:
    import py7zr
except ImportError:
    py7zr = None

try:
    import trimesh
except ImportError:
    trimesh = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


# ============================================================
# App settings
# ============================================================

APP_NAME = "fileconverter"

APP_FOLDER = Path.home() / ".fileconverter"
LIBRARY_FOLDER = APP_FOLDER / "library"
DATABASE_FILE = APP_FOLDER / "fileconverter.db"

APP_FOLDER.mkdir(exist_ok=True)
LIBRARY_FOLDER.mkdir(exist_ok=True)

MAX_BATCH_FILES = 100
FREE_UPLOAD_LIMIT_TOTAL = 10
PAID_MONTHLY_LIMIT = 1000
FREE_LIBRARY_DAYS = 30

PLAN_FREE = "Free"
PLAN_PRO = "Pro"
PLAN_ONE_TIME = "One-Time"
PLAN_MAX = "Max"

PLAN_DETAILS = {
    PLAN_FREE: {
        "price": "$0",
        "limit": "10 uploads total",
        "library": "Files stored in library are automatically deleted after 30 days."
    },
    PLAN_PRO: {
        "price": "$10/month",
        "limit": "1,000 uploads/month",
        "library": "Library files do not expire while subscribed."
    },
    PLAN_ONE_TIME: {
        "price": "$100 one-time",
        "limit": "1,000 uploads/month",
        "library": "Library files do not expire while subscribed."
    },
    PLAN_MAX: {
        "price": "$249.99 one-time + $20/month",
        "limit": "Unlimited uploads",
        "library": "Library files do not expire while subscribed."
    }
}


# ============================================================
# Supported formats
# ============================================================

IMAGE_FORMATS = {
    ".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff", ".tif", ".webp"
}

TEXT_FORMATS = {
    ".txt", ".docx", ".pdf"
}

SPREADSHEET_FORMATS = {
    ".csv", ".xlsx"
}

AUDIO_VIDEO_FORMATS = {
    ".mp3", ".wav", ".aac", ".flac",
    ".mp4", ".mov", ".avi", ".mkv", ".wmv"
}

MODEL_FORMATS = {
    ".glb", ".gltf", ".obj", ".stl", ".ply"
}

DISK_IMAGE_FORMATS = {
    ".img", ".iso"
}

BLOCKED_SYSTEM_FORMATS = {
    ".exe", ".dll", ".msi", ".bat", ".cmd", ".sh", ".app"
}

OUTPUT_FORMATS = [
    "txt", "pdf", "docx", "csv", "xlsx",
    "jpg", "jpeg", "png", "bmp", "gif", "tiff", "webp",
    "mp3", "wav", "aac", "flac",
    "mp4", "mov", "avi", "mkv", "wmv",
    "zip", "tar", "gz", "7z",
    "glb", "gltf", "obj", "stl", "ply"
]


# ============================================================
# Policy text
# ============================================================

POLICY_TEXT = f"""
FILECONVERTER USER AGREEMENT

By creating an account or logging into this app, you agree to the following:

1. File Conversion
fileconverter allows users to select files from their own device and convert them into supported formats.

2. Permission To Save Files
Before converted files are saved or downloaded onto your device, fileconverter will ask for permission.
You may choose where converted files are saved.

3. In-App Library
Users may store converted files inside their fileconverter library.
Library files are stored locally on this device, inside the fileconverter app folder.

4. User Deletion Control
fileconverter will not delete user files manually or randomly.
Users may delete one file or select multiple files to delete from their library.

5. 30-Day Free Plan Storage Policy
Free plan users are warned that library files may be automatically deleted after {FREE_LIBRARY_DAYS} days.
This keeps free storage limited and prevents the app folder from growing forever.

6. Subscribed User Storage Policy
Subscribed users do not have the 30-day automatic library deletion rule applied while subscribed.
Subscribed plans include Pro, One-Time, and Max.

7. Subscription Upload Limits
Free plan: 10 uploads total.
Pro plan: $10/month and 1,000 uploads per month.
One-Time plan: $100 one-time purchase and 1,000 uploads per month.
Max plan: $249.99 one-time fee plus $20/month and unlimited uploads.

8. Payment Notice
This demo version does not process real payments.
Upgrade buttons simulate subscription access for testing.
A real app would need Stripe, PayPal, Apple Pay, Google Pay, or another secure payment provider.

9. File Responsibility
You are responsible for only converting files you own or have permission to use.
Do not upload or convert illegal, stolen, copyrighted, or harmful files.

10. Security Notice
This is a local demo app. User accounts are stored locally on this computer.
Do not use your real banking, school, work, or personal password.
"""


# ============================================================
# Database helpers
# ============================================================

def connect_database():
    return sqlite3.connect(DATABASE_FILE)


def get_current_month():
    return datetime.now().strftime("%Y-%m")


def now_string():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def parse_datetime(value):
    return datetime.strptime(value, "%Y-%m-%d %H:%M:%S")


def setup_database():
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            salt TEXT NOT NULL,
            plan TEXT NOT NULL,
            total_uploads INTEGER NOT NULL,
            monthly_uploads INTEGER NOT NULL,
            current_month TEXT NOT NULL,
            accepted_policy INTEGER NOT NULL,
            created_at TEXT NOT NULL
        )
    """)

    cursor.execute("""
        CREATE TABLE IF NOT EXISTS library_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            original_name TEXT NOT NULL,
            stored_name TEXT NOT NULL,
            stored_path TEXT NOT NULL,
            output_format TEXT NOT NULL,
            file_size INTEGER NOT NULL,
            created_at TEXT NOT NULL,
            expires_at TEXT,
            FOREIGN KEY(user_id) REFERENCES users(id)
        )
    """)

    connection.commit()
    connection.close()


def hash_password(password, salt=None):
    if salt is None:
        salt = secrets.token_hex(16)

    password_hash = hashlib.sha256((salt + password).encode("utf-8")).hexdigest()
    return password_hash, salt


def create_user(username, password):
    setup_database()

    username = username.strip()

    if not username:
        raise ValueError("Username cannot be empty.")

    if len(password) < 4:
        raise ValueError("Password must be at least 4 characters.")

    password_hash, salt = hash_password(password)

    connection = connect_database()
    cursor = connection.cursor()

    try:
        cursor.execute("""
            INSERT INTO users (
                username, password_hash, salt, plan,
                total_uploads, monthly_uploads, current_month,
                accepted_policy, created_at
            )
            VALUES (?, ?, ?, ?, 0, 0, ?, 0, ?)
        """, (
            username,
            password_hash,
            salt,
            PLAN_FREE,
            get_current_month(),
            now_string()
        ))

        connection.commit()

    except sqlite3.IntegrityError:
        raise ValueError("That username already exists.")

    finally:
        connection.close()


def login_user(username, password):
    setup_database()

    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        SELECT id, username, password_hash, salt, plan,
               total_uploads, monthly_uploads, current_month,
               accepted_policy, created_at
        FROM users
        WHERE username = ?
    """, (username.strip(),))

    row = cursor.fetchone()
    connection.close()

    if row is None:
        raise ValueError("Account not found.")

    user_id, username, stored_hash, salt, plan, total_uploads, monthly_uploads, current_month, accepted_policy, created_at = row

    attempted_hash, _ = hash_password(password, salt)

    if attempted_hash != stored_hash:
        raise ValueError("Incorrect password.")

    if current_month != get_current_month():
        reset_monthly_uploads(user_id)
        monthly_uploads = 0
        current_month = get_current_month()

    return {
        "id": user_id,
        "username": username,
        "plan": plan,
        "total_uploads": total_uploads,
        "monthly_uploads": monthly_uploads,
        "current_month": current_month,
        "accepted_policy": bool(accepted_policy),
        "created_at": created_at
    }


def get_user_by_id(user_id):
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        SELECT id, username, plan, total_uploads, monthly_uploads,
               current_month, accepted_policy, created_at
        FROM users
        WHERE id = ?
    """, (user_id,))

    row = cursor.fetchone()
    connection.close()

    if row is None:
        return None

    user_id, username, plan, total_uploads, monthly_uploads, current_month, accepted_policy, created_at = row

    if current_month != get_current_month():
        reset_monthly_uploads(user_id)
        monthly_uploads = 0
        current_month = get_current_month()

    return {
        "id": user_id,
        "username": username,
        "plan": plan,
        "total_uploads": total_uploads,
        "monthly_uploads": monthly_uploads,
        "current_month": current_month,
        "accepted_policy": bool(accepted_policy),
        "created_at": created_at
    }


def accept_policy(user_id):
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        UPDATE users
        SET accepted_policy = 1
        WHERE id = ?
    """, (user_id,))

    connection.commit()
    connection.close()


def update_plan(user_id, new_plan):
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        UPDATE users
        SET plan = ?
        WHERE id = ?
    """, (new_plan, user_id))

    connection.commit()
    connection.close()


def reset_monthly_uploads(user_id):
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        UPDATE users
        SET monthly_uploads = 0,
            current_month = ?
        WHERE id = ?
    """, (get_current_month(), user_id))

    connection.commit()
    connection.close()


def add_upload_count(user_id, amount):
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        UPDATE users
        SET total_uploads = total_uploads + ?,
            monthly_uploads = monthly_uploads + ?
        WHERE id = ?
    """, (amount, amount, user_id))

    connection.commit()
    connection.close()


def get_remaining_uploads(user):
    plan = user["plan"]

    if plan == PLAN_FREE:
        remaining = FREE_UPLOAD_LIMIT_TOTAL - user["total_uploads"]
        return max(0, remaining)

    if plan in [PLAN_PRO, PLAN_ONE_TIME]:
        remaining = PAID_MONTHLY_LIMIT - user["monthly_uploads"]
        return max(0, remaining)

    if plan == PLAN_MAX:
        return "Unlimited"

    return 0


def can_convert_amount(user, file_count):
    plan = user["plan"]

    if plan == PLAN_MAX:
        return True, "Unlimited uploads available."

    if plan == PLAN_FREE:
        remaining = FREE_UPLOAD_LIMIT_TOTAL - user["total_uploads"]

        if file_count > remaining:
            return False, (
                f"Free plan limit reached.\n\n"
                f"You have {remaining} free upload(s) left.\n"
                f"You are trying to convert {file_count} file(s).\n\n"
                f"Upgrade to Pro, One-Time, or Max to continue."
            )

    if plan in [PLAN_PRO, PLAN_ONE_TIME]:
        remaining = PAID_MONTHLY_LIMIT - user["monthly_uploads"]

        if file_count > remaining:
            return False, (
                f"Monthly upload limit reached.\n\n"
                f"You have {remaining} upload(s) left this month.\n"
                f"You are trying to convert {file_count} file(s).\n\n"
                f"Upgrade to Max for unlimited uploads."
            )

    return True, "Allowed."


# ============================================================
# Library database helpers
# ============================================================

def get_user_library_folder(user):
    safe_username = "".join(c for c in user["username"] if c.isalnum() or c in ("_", "-")).strip()
    if not safe_username:
        safe_username = f"user_{user['id']}"

    folder = LIBRARY_FOLDER / f"{user['id']}_{safe_username}"
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def unique_file_path(folder, file_name):
    folder = Path(folder)
    original = Path(file_name)

    stem = original.stem
    suffix = original.suffix

    candidate = folder / file_name
    counter = 1

    while candidate.exists():
        candidate = folder / f"{stem}_{counter}{suffix}"
        counter += 1

    return candidate


def add_file_to_library_record(user, original_name, stored_path, output_format):
    stored_path = Path(stored_path)
    file_size = stored_path.stat().st_size if stored_path.exists() else 0
    created_at = now_string()

    expires_at = None

    if user["plan"] == PLAN_FREE:
        expires_at = (datetime.now() + timedelta(days=FREE_LIBRARY_DAYS)).strftime("%Y-%m-%d %H:%M:%S")

    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        INSERT INTO library_files (
            user_id, original_name, stored_name, stored_path,
            output_format, file_size, created_at, expires_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        user["id"],
        original_name,
        stored_path.name,
        str(stored_path),
        output_format,
        file_size,
        created_at,
        expires_at
    ))

    connection.commit()
    connection.close()


def get_library_files(user_id):
    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        SELECT id, original_name, stored_name, stored_path,
               output_format, file_size, created_at, expires_at
        FROM library_files
        WHERE user_id = ?
        ORDER BY id DESC
    """, (user_id,))

    rows = cursor.fetchall()
    connection.close()

    files = []

    for row in rows:
        file_id, original_name, stored_name, stored_path, output_format, file_size, created_at, expires_at = row

        files.append({
            "id": file_id,
            "original_name": original_name,
            "stored_name": stored_name,
            "stored_path": stored_path,
            "output_format": output_format,
            "file_size": file_size,
            "created_at": created_at,
            "expires_at": expires_at
        })

    return files


def delete_library_files_by_ids(user_id, file_ids):
    if not file_ids:
        return 0

    connection = connect_database()
    cursor = connection.cursor()

    deleted_count = 0

    for file_id in file_ids:
        cursor.execute("""
            SELECT stored_path
            FROM library_files
            WHERE id = ? AND user_id = ?
        """, (file_id, user_id))

        row = cursor.fetchone()

        if row is None:
            continue

        stored_path = Path(row[0])

        try:
            if stored_path.exists():
                stored_path.unlink()
        except Exception:
            pass

        cursor.execute("""
            DELETE FROM library_files
            WHERE id = ? AND user_id = ?
        """, (file_id, user_id))

        deleted_count += 1

    connection.commit()
    connection.close()

    return deleted_count


def auto_delete_expired_free_library_files(user):
    if user["plan"] != PLAN_FREE:
        return 0

    connection = connect_database()
    cursor = connection.cursor()

    cursor.execute("""
        SELECT id, stored_path
        FROM library_files
        WHERE user_id = ?
          AND expires_at IS NOT NULL
          AND expires_at < ?
    """, (user["id"], now_string()))

    rows = cursor.fetchall()

    deleted_count = 0

    for file_id, stored_path in rows:
        path = Path(stored_path)

        try:
            if path.exists():
                path.unlink()
        except Exception:
            pass

        cursor.execute("""
            DELETE FROM library_files
            WHERE id = ? AND user_id = ?
        """, (file_id, user["id"]))

        deleted_count += 1

    connection.commit()
    connection.close()

    return deleted_count


# ============================================================
# File conversion helpers
# ============================================================

def get_extension(file_path):
    return Path(file_path).suffix.lower()


def make_output_path(input_path, output_format, output_folder=None):
    input_path = Path(input_path)
    output_name = input_path.stem + "_converted." + output_format.lower()

    if output_folder:
        return str(unique_file_path(output_folder, output_name))

    return str(unique_file_path(input_path.parent, output_name))


def check_blocked(input_ext):
    if input_ext in BLOCKED_SYSTEM_FORMATS:
        raise ValueError(
            "Executable and system files cannot be converted into normal formats.\n\n"
            "Files like EXE, DLL, MSI, BAT, CMD, SH, and APP contain program data."
        )


def ffmpeg_available():
    try:
        subprocess.run(
            ["ffmpeg", "-version"],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            check=False
        )
        return True

    except FileNotFoundError:
        return False


def save_pil_image(img, output_path, output_format):
    output_format = output_format.lower()

    if output_format in ["jpg", "jpeg"]:
        img = img.convert("RGB")
        img.save(output_path, "JPEG", quality=95)

    elif output_format == "png":
        img.save(output_path, "PNG")

    elif output_format == "webp":
        img.save(output_path, "WEBP", quality=95)

    elif output_format == "bmp":
        img.save(output_path, "BMP")

    elif output_format == "gif":
        img.save(output_path, "GIF")

    elif output_format in ["tiff", "tif"]:
        img.save(output_path, "TIFF")

    else:
        raise ValueError("Unsupported image output format.")

    return output_path


def create_text_preview_image(title, lines, output_path, output_format):
    if Image is None or ImageDraw is None or ImageFont is None:
        raise ImportError("Pillow is missing. Install it with: python3 -m pip install pillow")

    width = 1200
    height = 900

    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    try:
        title_font = ImageFont.truetype("Arial.ttf", 42)
        body_font = ImageFont.truetype("Arial.ttf", 26)
    except Exception:
        title_font = ImageFont.load_default()
        body_font = ImageFont.load_default()

    x = 60
    y = 60

    draw.text((x, y), title, fill="black", font=title_font)
    y += 80

    for line in lines:
        if y > height - 70:
            draw.text((x, y), "...", fill="black", font=body_font)
            break

        line = str(line)

        while len(line) > 85:
            draw.text((x, y), line[:85], fill="black", font=body_font)
            line = line[85:]
            y += 36

            if y > height - 70:
                break

        if y > height - 70:
            break

        draw.text((x, y), line, fill="black", font=body_font)
        y += 36

    return save_pil_image(img, output_path, output_format)


# ============================================================
# Image conversion
# ============================================================

def convert_image(input_path, output_format, output_folder=None):
    if Image is None:
        raise ImportError("Pillow is missing. Install it with: python3 -m pip install pillow")

    output_path = make_output_path(input_path, output_format, output_folder)

    with Image.open(input_path) as img:
        return save_pil_image(img, output_path, output_format)


# ============================================================
# Document/text conversion
# ============================================================

def read_txt(input_path):
    with open(input_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()


def read_docx(input_path):
    if Document is None:
        raise ImportError("python-docx is missing. Install it with: python3 -m pip install python-docx")

    doc = Document(input_path)
    lines = []

    for paragraph in doc.paragraphs:
        lines.append(paragraph.text)

    return "\n".join(lines)


def read_pdf_text(input_path):
    if PdfReader is None:
        raise ImportError("pypdf is missing. Install it with: python3 -m pip install pypdf")

    reader = PdfReader(input_path)
    text = []

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text.append(page_text)

    return "\n\n".join(text)


def get_text_from_file(input_path):
    ext = get_extension(input_path)

    if ext == ".txt":
        return read_txt(input_path)

    if ext == ".docx":
        return read_docx(input_path)

    if ext == ".pdf":
        return read_pdf_text(input_path)

    raise ValueError("This document type can only be converted if it is TXT, DOCX, or PDF.")


def save_text_as_txt(text, output_path):
    with open(output_path, "w", encoding="utf-8") as file:
        file.write(text)


def save_text_as_docx(text, output_path):
    if Document is None:
        raise ImportError("python-docx is missing. Install it with: python3 -m pip install python-docx")

    doc = Document()

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(output_path)


def save_text_as_pdf(text, output_path):
    if canvas is None or letter is None:
        raise ImportError("reportlab is missing. Install it with: python3 -m pip install reportlab")

    pdf = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter

    x = 50
    y = height - 50
    line_height = 14

    for paragraph in text.split("\n"):
        paragraph = str(paragraph)

        while len(paragraph) > 90:
            part = paragraph[:90]
            pdf.drawString(x, y, part)
            paragraph = paragraph[90:]
            y -= line_height

            if y < 50:
                pdf.showPage()
                y = height - 50

        pdf.drawString(x, y, paragraph)
        y -= line_height

        if y < 50:
            pdf.showPage()
            y = height - 50

    pdf.save()


def convert_document_text(input_path, output_format, output_folder=None):
    output_path = make_output_path(input_path, output_format, output_folder)
    text = get_text_from_file(input_path)

    if output_format == "txt":
        save_text_as_txt(text, output_path)

    elif output_format == "docx":
        save_text_as_docx(text, output_path)

    elif output_format == "pdf":
        save_text_as_pdf(text, output_path)

    else:
        raise ValueError("Documents can only be converted to TXT, DOCX, or PDF.")

    return output_path


# ============================================================
# Convert many file types to image
# ============================================================

def convert_pdf_to_image(input_path, output_format, output_folder=None):
    if fitz is None:
        raise ImportError("PyMuPDF is missing. Install it with: python3 -m pip install pymupdf")

    if Image is None:
        raise ImportError("Pillow is missing. Install it with: python3 -m pip install pillow")

    output_path = make_output_path(input_path, output_format, output_folder)

    pdf = fitz.open(input_path)

    if len(pdf) == 0:
        raise ValueError("This PDF has no pages.")

    page = pdf[0]
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    pdf.close()

    return save_pil_image(img, output_path, output_format)


def convert_text_like_file_to_image(input_path, output_format, output_folder=None):
    output_path = make_output_path(input_path, output_format, output_folder)
    ext = get_extension(input_path)

    try:
        if ext in [".txt", ".csv"]:
            with open(input_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()

            lines = text.splitlines()[:40]

        elif ext == ".docx":
            text = read_docx(input_path)
            lines = text.splitlines()[:40]

        elif ext == ".xlsx":
            if pd is None:
                raise ImportError("pandas/openpyxl are missing.")

            df = pd.read_excel(input_path)
            lines = df.head(25).to_string(index=False).splitlines()

        else:
            lines = ["This file type cannot be shown directly as text."]

    except Exception as error:
        lines = [f"Could not read file content: {error}"]

    return create_text_preview_image(
        title=f"Preview: {Path(input_path).name}",
        lines=lines,
        output_path=output_path,
        output_format=output_format
    )


def convert_video_to_image(input_path, output_format, output_folder=None):
    if not ffmpeg_available():
        return create_unknown_file_preview(input_path, output_format, output_folder)

    output_path = make_output_path(input_path, output_format, output_folder)

    command = [
        "ffmpeg",
        "-y",
        "-i", input_path,
        "-frames:v", "1",
        output_path
    ]

    result = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )

    if result.returncode != 0:
        return create_unknown_file_preview(input_path, output_format, output_folder)

    return output_path


def create_unknown_file_preview(input_path, output_format, output_folder=None):
    output_path = make_output_path(input_path, output_format, output_folder)
    path = Path(input_path)

    try:
        size = path.stat().st_size
    except Exception:
        size = 0

    size_mb = size / (1024 * 1024)

    lines = [
        f"File name: {path.name}",
        f"File extension: {path.suffix if path.suffix else 'No extension'}",
        f"File size: {size_mb:.2f} MB",
        "",
        "This file type cannot be truly displayed as an image.",
        "A preview card image was created instead.",
        "",
        "This is useful for unknown files, system files, archives,",
        "installers, disk images, and other non-visual formats."
    ]

    return create_text_preview_image(
        title="File Preview Card",
        lines=lines,
        output_path=output_path,
        output_format=output_format
    )


def convert_any_file_to_image(input_path, output_format, output_folder=None):
    input_ext = get_extension(input_path)

    if output_format not in ["png", "jpg", "jpeg", "webp"]:
        raise ValueError("This feature only supports PNG, JPG, JPEG, and WEBP.")

    if input_ext in IMAGE_FORMATS:
        return convert_image(input_path, output_format, output_folder)

    if input_ext == ".pdf":
        return convert_pdf_to_image(input_path, output_format, output_folder)

    if input_ext in [".txt", ".docx", ".csv", ".xlsx"]:
        return convert_text_like_file_to_image(input_path, output_format, output_folder)

    if input_ext in AUDIO_VIDEO_FORMATS:
        return convert_video_to_image(input_path, output_format, output_folder)

    return create_unknown_file_preview(input_path, output_format, output_folder)


# ============================================================
# Spreadsheet conversion
# ============================================================

def convert_spreadsheet(input_path, output_format, output_folder=None):
    if pd is None:
        raise ImportError("pandas/openpyxl are missing. Install them with: python3 -m pip install pandas openpyxl")

    input_ext = get_extension(input_path)
    output_path = make_output_path(input_path, output_format, output_folder)

    if input_ext == ".csv" and output_format == "xlsx":
        df = pd.read_csv(input_path)
        df.to_excel(output_path, index=False)
        return output_path

    if input_ext == ".xlsx" and output_format == "csv":
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False)
        return output_path

    raise ValueError("Spreadsheet conversion only supports CSV and XLSX.")


# ============================================================
# Audio/video conversion
# ============================================================

def convert_audio_video(input_path, output_format, output_folder=None):
    if not ffmpeg_available():
        raise EnvironmentError(
            "FFmpeg is not installed.\n\n"
            "On Mac, install it with:\n"
            "brew install ffmpeg"
        )

    output_path = make_output_path(input_path, output_format, output_folder)

    command = [
        "ffmpeg",
        "-y",
        "-i", input_path,
        output_path
    ]

    result = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True
    )

    if result.returncode != 0:
        raise RuntimeError("FFmpeg failed:\n\n" + result.stderr)

    return output_path


# ============================================================
# 3D model conversion
# ============================================================

def convert_model(input_path, output_format, output_folder=None):
    if trimesh is None:
        raise ImportError("trimesh is missing. Install it with: python3 -m pip install trimesh")

    output_ext = "." + output_format

    if output_ext not in MODEL_FORMATS:
        raise ValueError("3D models can only be converted to GLB, GLTF, OBJ, STL, or PLY.")

    output_path = make_output_path(input_path, output_format, output_folder)

    model = trimesh.load(input_path, force="scene")

    if model is None:
        raise ValueError("Could not load this 3D model.")

    model.export(output_path)

    return output_path


# ============================================================
# Archive conversion
# ============================================================

def convert_to_zip(input_path, output_folder=None):
    output_path = make_output_path(input_path, "zip", output_folder)

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zip_file:
        path = Path(input_path)

        if path.is_file():
            zip_file.write(path, arcname=path.name)
        else:
            for item in path.rglob("*"):
                zip_file.write(item, arcname=item.relative_to(path.parent))

    return output_path


def convert_to_tar(input_path, output_folder=None):
    output_path = make_output_path(input_path, "tar", output_folder)

    with tarfile.open(output_path, "w") as tar:
        tar.add(input_path, arcname=Path(input_path).name)

    return output_path


def convert_to_gz(input_path, output_folder=None):
    path = Path(input_path)

    if not path.is_file():
        raise ValueError("GZ compression only works on one file, not a folder.")

    output_path = make_output_path(input_path, "gz", output_folder)

    with open(input_path, "rb") as source:
        with gzip.open(output_path, "wb") as target:
            shutil.copyfileobj(source, target)

    return output_path


def convert_to_7z(input_path, output_folder=None):
    if py7zr is None:
        raise ImportError("py7zr is missing. Install it with: python3 -m pip install py7zr")

    output_path = make_output_path(input_path, "7z", output_folder)
    path = Path(input_path)

    with py7zr.SevenZipFile(output_path, "w") as archive:
        if path.is_file():
            archive.write(path, arcname=path.name)
        else:
            archive.writeall(path, arcname=path.name)

    return output_path


def convert_archive(input_path, output_format, output_folder=None):
    if output_format == "zip":
        return convert_to_zip(input_path, output_folder)

    if output_format == "tar":
        return convert_to_tar(input_path, output_folder)

    if output_format == "gz":
        return convert_to_gz(input_path, output_folder)

    if output_format == "7z":
        return convert_to_7z(input_path, output_folder)

    raise ValueError("Archive output must be ZIP, TAR, GZ, or 7Z.")


# ============================================================
# Main conversion controller
# ============================================================

def convert_file(input_path, output_format, output_folder=None):
    input_ext = get_extension(input_path)
    output_format = output_format.lower()
    output_ext = "." + output_format

    if output_format in ["png", "jpg", "jpeg", "webp"]:
        return convert_any_file_to_image(input_path, output_format, output_folder)

    check_blocked(input_ext)

    if input_ext in IMAGE_FORMATS:
        if output_ext not in IMAGE_FORMATS:
            raise ValueError("Image files can only be converted to another image format.")

        return convert_image(input_path, output_format, output_folder)

    if input_ext in TEXT_FORMATS:
        if output_format not in ["txt", "docx", "pdf"]:
            raise ValueError("Documents can only be converted to TXT, DOCX, or PDF.")

        return convert_document_text(input_path, output_format, output_folder)

    if input_ext in SPREADSHEET_FORMATS:
        if output_format not in ["csv", "xlsx"]:
            raise ValueError("Spreadsheets can only be converted between CSV and XLSX.")

        return convert_spreadsheet(input_path, output_format, output_folder)

    if input_ext in AUDIO_VIDEO_FORMATS:
        if output_ext not in AUDIO_VIDEO_FORMATS:
            raise ValueError("Audio/video files can only be converted to another audio/video format.")

        return convert_audio_video(input_path, output_format, output_folder)

    if input_ext in MODEL_FORMATS:
        if output_ext not in MODEL_FORMATS:
            raise ValueError("3D model files can only be converted to GLB, GLTF, OBJ, STL, or PLY.")

        return convert_model(input_path, output_format, output_folder)

    if input_ext in DISK_IMAGE_FORMATS:
        if output_format not in ["zip", "tar", "gz", "7z"]:
            raise ValueError(
                "IMG and ISO files are disk-image files.\n\n"
                "This app can compress them to ZIP, TAR, GZ, or 7Z, "
                "but it does not rewrite their internal disk-image format."
            )

        return convert_archive(input_path, output_format, output_folder)

    if output_format in ["zip", "tar", "gz", "7z"]:
        return convert_archive(input_path, output_format, output_folder)

    raise ValueError("Unsupported conversion.")


# ============================================================
# Login window
# ============================================================

class LoginWindow:
    def __init__(self, root):
        self.root = root
        self.root.title("fileconverter - Login")
        self.root.geometry("500x420")
        self.root.resizable(False, False)

        self.current_user = None
        self.build_ui()

    def build_ui(self):
        title = tk.Label(
            self.root,
            text="fileconverter",
            font=("Arial", 28, "bold")
        )
        title.pack(pady=20)

        subtitle = tk.Label(
            self.root,
            text="Sign up or log in to use your converter library.",
            font=("Arial", 12)
        )
        subtitle.pack(pady=5)

        form = tk.Frame(self.root)
        form.pack(pady=25)

        username_label = tk.Label(form, text="Username:", font=("Arial", 12))
        username_label.grid(row=0, column=0, sticky="e", padx=5, pady=8)

        self.username_entry = tk.Entry(form, width=28, font=("Arial", 12))
        self.username_entry.grid(row=0, column=1, padx=5, pady=8)

        password_label = tk.Label(form, text="Password:", font=("Arial", 12))
        password_label.grid(row=1, column=0, sticky="e", padx=5, pady=8)

        self.password_entry = tk.Entry(form, width=28, font=("Arial", 12), show="*")
        self.password_entry.grid(row=1, column=1, padx=5, pady=8)

        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=15)

        login_button = tk.Button(
            button_frame,
            text="Login",
            command=self.login,
            width=16,
            height=2
        )
        login_button.grid(row=0, column=0, padx=10)

        signup_button = tk.Button(
            button_frame,
            text="Sign Up",
            command=self.signup,
            width=16,
            height=2
        )
        signup_button.grid(row=0, column=1, padx=10)

        warning = tk.Label(
            self.root,
            text="Demo app notice: Do not use your real banking, school, or work password.",
            fg="gray25",
            font=("Arial", 10),
            wraplength=420
        )
        warning.pack(pady=20)

    def get_form_values(self):
        username = self.username_entry.get().strip()
        password = self.password_entry.get()

        if not username:
            raise ValueError("Enter a username.")

        if not password:
            raise ValueError("Enter a password.")

        return username, password

    def signup(self):
        try:
            username, password = self.get_form_values()
            create_user(username, password)
            messagebox.showinfo("Account Created", "Your account was created. You can now log in.")

        except Exception as error:
            messagebox.showerror("Sign Up Failed", str(error))

    def login(self):
        try:
            username, password = self.get_form_values()
            user = login_user(username, password)

            self.current_user = user

            self.root.destroy()
            launch_main_app(user)

        except Exception as error:
            messagebox.showerror("Login Failed", str(error))


# ============================================================
# Policy agreement window
# ============================================================

class PolicyWindow:
    def __init__(self, parent, user):
        self.parent = parent
        self.user = user
        self.accepted = False

        self.window = tk.Toplevel(parent)
        self.window.title("fileconverter - User Agreement")
        self.window.geometry("760x620")
        self.window.resizable(False, False)
        self.window.grab_set()

        self.build_ui()

    def build_ui(self):
        title = tk.Label(
            self.window,
            text="User Agreement and Storage Policy",
            font=("Arial", 20, "bold")
        )
        title.pack(pady=10)

        text_frame = tk.Frame(self.window)
        text_frame.pack(pady=10)

        self.text_box = tk.Text(
            text_frame,
            width=88,
            height=25,
            wrap="word"
        )
        self.text_box.grid(row=0, column=0)

        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.grid(row=0, column=1, sticky="ns")

        self.text_box.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.text_box.yview)

        self.text_box.insert("1.0", POLICY_TEXT)
        self.text_box.config(state="disabled")

        self.agree_var = tk.BooleanVar(value=False)

        agree_check = tk.Checkbutton(
            self.window,
            text="I have read and agree to this policy.",
            variable=self.agree_var,
            font=("Arial", 11)
        )
        agree_check.pack(pady=10)

        buttons = tk.Frame(self.window)
        buttons.pack(pady=10)

        accept_button = tk.Button(
            buttons,
            text="Accept and Continue",
            command=self.accept,
            width=22,
            height=2
        )
        accept_button.grid(row=0, column=0, padx=10)

        decline_button = tk.Button(
            buttons,
            text="Decline and Exit",
            command=self.decline,
            width=18,
            height=2
        )
        decline_button.grid(row=0, column=1, padx=10)

    def accept(self):
        if not self.agree_var.get():
            messagebox.showwarning("Agreement Required", "You must check the agreement box to continue.")
            return

        accept_policy(self.user["id"])
        self.accepted = True
        self.window.destroy()

    def decline(self):
        messagebox.showinfo("Agreement Declined", "You must accept the policy to use fileconverter.")
        self.window.destroy()
        self.parent.destroy()


# ============================================================
# Main app
# ============================================================

class FileConverterApp:
    def __init__(self, root, user):
        self.root = root
        self.user = user

        self.root.title(APP_NAME)
        self.root.geometry("1050x760")
        self.root.resizable(False, False)

        self.selected_files = []
        self.output_format = tk.StringVar(value="png")

        self.save_to_device_var = tk.BooleanVar(value=True)
        self.store_library_var = tk.BooleanVar(value=True)

        self.library_items = []

        self.require_policy_if_needed()
        self.run_free_cleanup()
        self.build_ui()
        self.refresh_all()

    def require_policy_if_needed(self):
        fresh_user = get_user_by_id(self.user["id"])
        self.user = fresh_user

        if not self.user["accepted_policy"]:
            policy = PolicyWindow(self.root, self.user)
            self.root.wait_window(policy.window)

            fresh_user = get_user_by_id(self.user["id"])

            if not fresh_user or not fresh_user["accepted_policy"]:
                return

            self.user = fresh_user

    def run_free_cleanup(self):
        deleted = auto_delete_expired_free_library_files(self.user)

        if deleted > 0:
            messagebox.showwarning(
                "Expired Free Library Files Deleted",
                f"{deleted} file(s) were automatically deleted because they were stored for over "
                f"{FREE_LIBRARY_DAYS} days on the Free plan.\n\n"
                "Subscribed users do not have this 30-day library expiration rule."
            )

    def build_ui(self):
        title = tk.Label(
            self.root,
            text="fileconverter",
            font=("Arial", 26, "bold")
        )
        title.pack(pady=8)

        self.account_label = tk.Label(
            self.root,
            text="",
            font=("Arial", 11),
            fg="gray20"
        )
        self.account_label.pack(pady=3)

        notebook = ttk.Notebook(self.root)
        notebook.pack(fill="both", expand=True, padx=15, pady=10)

        self.converter_tab = tk.Frame(notebook)
        self.library_tab = tk.Frame(notebook)
        self.policy_tab = tk.Frame(notebook)

        notebook.add(self.converter_tab, text="Convert Files")
        notebook.add(self.library_tab, text="My Library")
        notebook.add(self.policy_tab, text="Policy Agreement")

        self.build_converter_tab()
        self.build_library_tab()
        self.build_policy_tab()

    def build_converter_tab(self):
        plan_frame = tk.LabelFrame(
            self.converter_tab,
            text="Subscription Plans",
            font=("Arial", 11, "bold"),
            padx=10,
            pady=10
        )
        plan_frame.pack(pady=8)

        pro_button = tk.Button(
            plan_frame,
            text="Upgrade to Pro - $10/month",
            command=lambda: self.upgrade_plan(PLAN_PRO),
            width=28
        )
        pro_button.grid(row=0, column=0, padx=5, pady=5)

        one_time_button = tk.Button(
            plan_frame,
            text="Buy One-Time - $100",
            command=lambda: self.upgrade_plan(PLAN_ONE_TIME),
            width=28
        )
        one_time_button.grid(row=0, column=1, padx=5, pady=5)

        max_button = tk.Button(
            plan_frame,
            text="Upgrade to Max - $249.99 + $20/month",
            command=lambda: self.upgrade_plan(PLAN_MAX),
            width=35
        )
        max_button.grid(row=0, column=2, padx=5, pady=5)

        free_button = tk.Button(
            plan_frame,
            text="Reset to Free Demo",
            command=lambda: self.upgrade_plan(PLAN_FREE),
            width=22
        )
        free_button.grid(row=1, column=1, padx=5, pady=5)

        button_frame = tk.Frame(self.converter_tab)
        button_frame.pack(pady=8)

        choose_files_button = tk.Button(
            button_frame,
            text="Choose Files",
            command=self.choose_files,
            width=18
        )
        choose_files_button.grid(row=0, column=0, padx=5)

        choose_folder_button = tk.Button(
            button_frame,
            text="Choose Folder",
            command=self.choose_folder,
            width=18
        )
        choose_folder_button.grid(row=0, column=1, padx=5)

        clear_button = tk.Button(
            button_frame,
            text="Clear List",
            command=self.clear_files,
            width=18
        )
        clear_button.grid(row=0, column=2, padx=5)

        list_frame = tk.Frame(self.converter_tab)
        list_frame.pack(pady=8)

        self.file_listbox = tk.Listbox(
            list_frame,
            width=115,
            height=9
        )
        self.file_listbox.grid(row=0, column=0)

        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.grid(row=0, column=1, sticky="ns")

        self.file_listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.file_listbox.yview)

        self.count_label = tk.Label(
            self.converter_tab,
            text="0 files selected",
            font=("Arial", 10)
        )
        self.count_label.pack(pady=2)

        format_frame = tk.Frame(self.converter_tab)
        format_frame.pack(pady=8)

        format_label = tk.Label(
            format_frame,
            text="Convert to:",
            font=("Arial", 11)
        )
        format_label.grid(row=0, column=0, padx=5)

        format_box = ttk.Combobox(
            format_frame,
            textvariable=self.output_format,
            values=OUTPUT_FORMATS,
            width=15,
            state="readonly"
        )
        format_box.grid(row=0, column=1, padx=5)

        options_frame = tk.LabelFrame(
            self.converter_tab,
            text="Save Options",
            padx=10,
            pady=8
        )
        options_frame.pack(pady=8)

        save_device_check = tk.Checkbutton(
            options_frame,
            text="Ask permission and save/download converted files to my device",
            variable=self.save_to_device_var,
            font=("Arial", 10)
        )
        save_device_check.grid(row=0, column=0, sticky="w", padx=5, pady=2)

        store_library_check = tk.Checkbutton(
            options_frame,
            text="Also store converted files in my in-app library",
            variable=self.store_library_var,
            font=("Arial", 10)
        )
        store_library_check.grid(row=1, column=0, sticky="w", padx=5, pady=2)

        convert_button = tk.Button(
            self.converter_tab,
            text="Convert Selected Files",
            command=self.handle_convert,
            font=("Arial", 14, "bold"),
            width=26,
            height=2
        )
        convert_button.pack(pady=10)

        note = (
            "Free library warning: Free users get 10 uploads total, and library files expire after 30 days. "
            "Subscribed users do not have the 30-day library expiration rule."
        )

        note_label = tk.Label(
            self.converter_tab,
            text=note,
            font=("Arial", 10),
            fg="gray25",
            wraplength=850
        )
        note_label.pack(pady=5)

    def build_library_tab(self):
        title = tk.Label(
            self.library_tab,
            text="My Converted File Library",
            font=("Arial", 20, "bold")
        )
        title.pack(pady=10)

        warning = tk.Label(
            self.library_tab,
            text=(
                "You control library deletion. Select one or more files and click Delete Selected. "
                "Free plan files automatically expire after 30 days. Subscribed users do not have that expiration rule."
            ),
            font=("Arial", 10),
            fg="gray25",
            wraplength=900
        )
        warning.pack(pady=5)

        list_frame = tk.Frame(self.library_tab)
        list_frame.pack(pady=10)

        self.library_listbox = tk.Listbox(
            list_frame,
            width=125,
            height=18,
            selectmode=tk.EXTENDED
        )
        self.library_listbox.grid(row=0, column=0)

        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.grid(row=0, column=1, sticky="ns")

        self.library_listbox.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.library_listbox.yview)

        buttons = tk.Frame(self.library_tab)
        buttons.pack(pady=10)

        refresh_button = tk.Button(
            buttons,
            text="Refresh Library",
            command=self.refresh_library_display,
            width=18
        )
        refresh_button.grid(row=0, column=0, padx=5)

        open_button = tk.Button(
            buttons,
            text="Open Selected File Location",
            command=self.open_selected_file_location,
            width=24
        )
        open_button.grid(row=0, column=1, padx=5)

        delete_button = tk.Button(
            buttons,
            text="Delete Selected",
            command=self.delete_selected_library_files,
            width=18
        )
        delete_button.grid(row=0, column=2, padx=5)

        self.library_count_label = tk.Label(
            self.library_tab,
            text="0 files in library",
            font=("Arial", 10)
        )
        self.library_count_label.pack(pady=5)

    def build_policy_tab(self):
        title = tk.Label(
            self.policy_tab,
            text="Policy Agreement",
            font=("Arial", 20, "bold")
        )
        title.pack(pady=10)

        text_frame = tk.Frame(self.policy_tab)
        text_frame.pack(pady=10)

        policy_box = tk.Text(
            text_frame,
            width=110,
            height=25,
            wrap="word"
        )
        policy_box.grid(row=0, column=0)

        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.grid(row=0, column=1, sticky="ns")

        policy_box.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=policy_box.yview)

        policy_box.insert("1.0", POLICY_TEXT)
        policy_box.config(state="disabled")

    def refresh_all(self):
        self.user = get_user_by_id(self.user["id"])
        self.refresh_account_display()
        self.refresh_file_listbox()
        self.refresh_library_display()

    def refresh_account_display(self):
        self.user = get_user_by_id(self.user["id"])
        remaining = get_remaining_uploads(self.user)

        text = (
            f"Logged in as: {self.user['username']} | "
            f"Plan: {self.user['plan']} | "
            f"Total Uploads: {self.user['total_uploads']} | "
            f"Monthly Uploads: {self.user['monthly_uploads']} | "
            f"Remaining: {remaining}"
        )

        self.account_label.config(text=text)

    def upgrade_plan(self, plan):
        if plan == PLAN_FREE:
            confirm = messagebox.askyesno(
                "Reset to Free",
                "This will switch your demo account back to Free.\n\n"
                f"Free users are warned that library files older than {FREE_LIBRARY_DAYS} days may be automatically deleted.\n\n"
                "Continue?"
            )

            if not confirm:
                return

            update_plan(self.user["id"], PLAN_FREE)
            self.user = get_user_by_id(self.user["id"])
            self.run_free_cleanup()
            self.refresh_all()
            return

        details = PLAN_DETAILS[plan]

        message = (
            f"Plan: {plan}\n"
            f"Price: {details['price']}\n"
            f"Limit: {details['limit']}\n"
            f"Library: {details['library']}\n\n"
            "This is a demo upgrade screen.\n"
            "In a real app, this button would open Stripe, PayPal, or another checkout page.\n\n"
            "Activate this plan for testing?"
        )

        confirm = messagebox.askyesno("Upgrade Plan", message)

        if confirm:
            update_plan(self.user["id"], plan)
            self.user = get_user_by_id(self.user["id"])
            self.refresh_all()
            messagebox.showinfo("Plan Updated", f"Your demo plan is now: {plan}")

    def refresh_file_listbox(self):
        self.file_listbox.delete(0, tk.END)

        for file_path in self.selected_files:
            self.file_listbox.insert(tk.END, file_path)

        self.count_label.config(text=f"{len(self.selected_files)} files selected")

    def choose_files(self):
        files = filedialog.askopenfilenames(title="Choose up to 100 files")

        if not files:
            return

        combined_files = self.selected_files + list(files)

        seen = set()
        unique_files = []

        for file_path in combined_files:
            if file_path not in seen:
                seen.add(file_path)
                unique_files.append(file_path)

        if len(unique_files) > MAX_BATCH_FILES:
            messagebox.showwarning(
                "Too Many Files",
                f"You selected more than {MAX_BATCH_FILES} files.\n\n"
                f"Only the first {MAX_BATCH_FILES} files will be added."
            )

            unique_files = unique_files[:MAX_BATCH_FILES]

        self.selected_files = unique_files
        self.refresh_file_listbox()

    def choose_folder(self):
        folder_path = filedialog.askdirectory(title="Choose a folder")

        if not folder_path:
            return

        folder = Path(folder_path)
        files = [str(path) for path in folder.iterdir() if path.is_file()]

        if len(files) > MAX_BATCH_FILES:
            messagebox.showwarning(
                "Too Many Files",
                f"This folder has more than {MAX_BATCH_FILES} files.\n\n"
                f"Only the first {MAX_BATCH_FILES} files will be added."
            )

            files = files[:MAX_BATCH_FILES]

        self.selected_files = files
        self.refresh_file_listbox()

    def clear_files(self):
        self.selected_files = []
        self.refresh_file_listbox()

    def handle_convert(self):
        if not self.selected_files:
            messagebox.showwarning("No Files Selected", "Please choose at least one file first.")
            return

        if not self.save_to_device_var.get() and not self.store_library_var.get():
            messagebox.showwarning(
                "No Save Option Selected",
                "Choose at least one option:\n\n"
                "- Save/download to device\n"
                "- Store in app library"
            )
            return

        self.user = get_user_by_id(self.user["id"])
        allowed, reason = can_convert_amount(self.user, len(self.selected_files))

        if not allowed:
            messagebox.showwarning("Subscription Required", reason)
            return

        output_format = self.output_format.get().strip().lower()

        device_folder = None

        if self.save_to_device_var.get():
            permission = messagebox.askyesno(
                "Permission Required",
                "fileconverter wants permission to save/download the converted file(s) onto your device.\n\n"
                "Do you give permission to continue?"
            )

            if not permission:
                messagebox.showinfo(
                    "Permission Denied",
                    "Device download/save was cancelled because permission was not granted."
                )

                if not self.store_library_var.get():
                    return

            else:
                device_folder = filedialog.askdirectory(
                    title="Choose where to download/save your converted files"
                )

                if not device_folder:
                    messagebox.showinfo(
                        "No Folder Selected",
                        "Device download/save was cancelled because no folder was selected."
                    )

                    if not self.store_library_var.get():
                        return

        library_folder = get_user_library_folder(self.user) if self.store_library_var.get() else None

        successful = []
        failed = []

        for file_path in self.selected_files:
            try:
                if not os.path.exists(file_path):
                    raise FileNotFoundError("File does not exist.")

                primary_output_folder = None

                if device_folder:
                    primary_output_folder = device_folder
                elif library_folder:
                    primary_output_folder = library_folder

                output_path = convert_file(file_path, output_format, primary_output_folder)

                saved_paths = []

                if device_folder:
                    saved_paths.append(output_path)

                if library_folder:
                    if Path(output_path).parent == Path(library_folder):
                        library_path = Path(output_path)
                    else:
                        library_path = unique_file_path(library_folder, Path(output_path).name)
                        shutil.copy2(output_path, library_path)

                    add_file_to_library_record(
                        user=self.user,
                        original_name=Path(file_path).name,
                        stored_path=library_path,
                        output_format=output_format
                    )

                    saved_paths.append(str(library_path))

                successful.append({
                    "input": file_path,
                    "output": output_path,
                    "saved_paths": saved_paths
                })

            except Exception as error:
                failed.append((file_path, str(error)))

        if successful:
            add_upload_count(self.user["id"], len(successful))

        self.refresh_all()

        message = (
            f"Conversion finished.\n\n"
            f"Successful: {len(successful)}\n"
            f"Failed: {len(failed)}"
        )

        if device_folder:
            message += f"\n\nDownloaded/Saved to device folder:\n{device_folder}"

        if library_folder:
            message += f"\n\nStored in app library:\n{library_folder}"

        if successful:
            message += "\n\nConverted files:\n"

            for item in successful[:10]:
                message += f"- {Path(item['output']).name}\n"

            if len(successful) > 10:
                message += f"...and {len(successful) - 10} more.\n"

        if failed:
            message += "\n\nFailed files:\n"

            for failed_file, error in failed[:5]:
                message += f"- {Path(failed_file).name}: {error}\n"

            if len(failed) > 5:
                message += f"...and {len(failed) - 5} more failed.\n"

        if failed:
            messagebox.showwarning("Conversion Complete With Errors", message)
        else:
            messagebox.showinfo("Conversion Complete", message)

    def refresh_library_display(self):
        self.library_items = get_library_files(self.user["id"])
        self.library_listbox.delete(0, tk.END)

        for item in self.library_items:
            size_mb = item["file_size"] / (1024 * 1024)

            expires_text = "No expiration"

            if item["expires_at"]:
                expires_text = f"Expires: {item['expires_at']}"

            display_text = (
                f"{item['stored_name']} | "
                f"Original: {item['original_name']} | "
                f"Format: {item['output_format']} | "
                f"Size: {size_mb:.2f} MB | "
                f"Created: {item['created_at']} | "
                f"{expires_text}"
            )

            self.library_listbox.insert(tk.END, display_text)

        self.library_count_label.config(text=f"{len(self.library_items)} files in library")

    def get_selected_library_ids(self):
        selected_indexes = self.library_listbox.curselection()
        selected_ids = []

        for index in selected_indexes:
            if 0 <= index < len(self.library_items):
                selected_ids.append(self.library_items[index]["id"])

        return selected_ids

    def open_selected_file_location(self):
        selected_indexes = self.library_listbox.curselection()

        if not selected_indexes:
            messagebox.showwarning("No Selection", "Select one library file first.")
            return

        index = selected_indexes[0]
        item = self.library_items[index]
        path = Path(item["stored_path"])

        if not path.exists():
            messagebox.showerror("File Missing", "This file no longer exists on disk.")
            return

        try:
            if os.name == "nt":
                os.startfile(path.parent)
            elif sys_platform_is_mac():
                subprocess.run(["open", str(path.parent)], check=False)
            else:
                subprocess.run(["xdg-open", str(path.parent)], check=False)

        except Exception as error:
            messagebox.showerror("Open Failed", str(error))

    def delete_selected_library_files(self):
        selected_ids = self.get_selected_library_ids()

        if not selected_ids:
            messagebox.showwarning("No Files Selected", "Select one or more files to delete.")
            return

        confirm = messagebox.askyesno(
            "Confirm Delete",
            f"You selected {len(selected_ids)} file(s).\n\n"
            "This will delete them from your in-app library and remove the stored file copies from this device.\n\n"
            "fileconverter will never delete selected files unless you confirm.\n\n"
            "Delete selected files?"
        )

        if not confirm:
            return

        deleted_count = delete_library_files_by_ids(self.user["id"], selected_ids)

        self.refresh_library_display()

        messagebox.showinfo(
            "Files Deleted",
            f"{deleted_count} file(s) were deleted from your library."
        )


def sys_platform_is_mac():
    import sys
    return sys.platform == "darwin"


# ============================================================
# Launch helpers
# ============================================================

def launch_main_app(user):
    root = tk.Tk()
    app = FileConverterApp(root, user)
    root.mainloop()


def main():
    setup_database()

    root = tk.Tk()
    login = LoginWindow(root)
    root.mainloop()


if __name__ == "__main__":
    main()